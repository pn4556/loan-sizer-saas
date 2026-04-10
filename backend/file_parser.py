"""
Universal File Parser for Loan Applications
Supports: PDF, DOCX, TXT, EML, MSG, Images (OCR), CSV, Excel
"""

import os
import re
import io
from typing import Dict, List, Optional, Tuple, Union
from dataclasses import dataclass
from pathlib import Path
import logging

logger = logging.getLogger(__name__)

@dataclass
class ExtractionResult:
    """Result of file extraction"""
    success: bool
    text: str
    fields: Dict[str, any]
    confidence: float
    missing_fields: List[str]
    error_message: Optional[str] = None
    file_type: str = "unknown"


class UniversalFileParser:
    """Parser for multiple file formats"""
    
    SUPPORTED_FORMATS = {
        'pdf': ['.pdf'],
        'docx': ['.docx', '.doc'],
        'text': ['.txt', '.text'],
        'email': ['.eml', '.msg'],
        'image': ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif'],
        'spreadsheet': ['.csv', '.xlsx', '.xls']
    }
    
    def __init__(self):
        self.patterns = self._compile_patterns()
    
    def _compile_patterns(self) -> Dict:
        """Compile regex patterns for field extraction"""
        return {
            'units': [
                r'(\d+)\s*(?:unit|units|doors?)',
                r'number\s+of\s+units?[:\s]+(\d+)',
                r'total\s+units?[:\s]+(\d+)',
                r'(?:property|building)\s*:?\s*(\d+)\s*(?:unit|units)',
            ],
            'address': [
                r'(?:property\s+)?address[:\s]+([^\n,]+(?:street|st|avenue|ave|road|rd|drive|dr|boulevard|blvd|lane|ln|way|circle|cir|court|ct)[^\n,]*)',
                r'location[:\s]+([^\n,]+(?:street|st|avenue|ave)[^\n,]*)',
                r'property\s+located\s+(?:at|on)[:\s]+([^\n,]+)',
            ],
            'city': [
                r'city[:\s]+([a-zA-Z\s]+?)(?:,|\n|$|state)',
                r'[^,]+,\s*([a-zA-Z\s]+?),\s*[A-Za-z]{2}\s*\d{5}',
            ],
            'state': [
                r',\s*([A-Z]{2})\s*\d{5}',
                r'state[:\s]+([A-Z]{2})',
                r'state[:\s]+([A-Za-z\s]+?)(?:\n|$|\d)',
            ],
            'zip': [
                r'\b(\d{5}(?:-\d{4})?)\b',
                r'zip[:\s]+(\d{5})',
                r'postal\s+code[:\s]+(\d{5})',
            ],
            'estimated_value': [
                r'(?:est(?:imated)?\s+)?(?:value|appraised\s+value|market\s+value)[:\s]+\$?([\d,\.]+)',
                r'property\s+value[:\s]+\$?([\d,\.]+)',
                r'estimated\s+value[:\s]+\$?([\d,\.]+)',
            ],
            'purchase_price': [
                r'(?:purchase\s+price|purchase\s+amount|acquisition\s+price|cost)[:\s]+\$?([\d,\.]+)',
            ],
            'loan_amount': [
                r'(?:loan\s+amount|requested\s+loan|financing\s+amount|loan\s+request)[:\s]+\$?([\d,\.]+)',
                r'loan[:\s]+\$?([\d,\.]+)',
            ],
            'note_type': [
                r'(30|15|20|10)\s*(?:year|yr)?\s*(?:fixed|arm|variable)',
                r'loan\s+type[:\s]+(30\s*year\s*fixed|15\s*year\s*fixed|ARM|variable)',
                r'(?:term|note)[:\s]+(\d+)\s*(?:year|yr)',
            ],
            'points': [
                r'(?:points|origination|origination\s+points)[:\s]+(\d+\.?\d*)',
                r'(\d+\.?\d*)\s*(?:point|pts)',
            ],
            'credit_scores': [
                r'(?:credit\s+score|fico)[s:]?\s*(\d{3})[\s,]+(\d{3})[\s,]+(\d{3})',
                r'(?:experian|exp)[:\s]*(\d{3}).{0,50}(?:transunion|trans)[:\s]*(\d{3}).{0,50}(?:equifax|eq)[:\s]*(\d{3})',
                r'credit[:\s]+(\d{3})\D+(\d{3})\D+(\d{3})',
            ],
            'applicant_name': [
                r'(?:borrower|applicant|name|submitted\s+by)[:\s]+([A-Za-z\s\.]+?)(?:\n|email|$)',
                r'borrower[:\s]+([A-Za-z\s\.]+)',
            ],
            'applicant_email': [
                r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
            ],
            'property_type': [
                r'property\s+type[:\s]+(multifamily|commercial|mixed-use|mixed\s+use|retail|office|industrial|single\s+family)',
                r'(multifamily|mixed-use|commercial)\s+property',
            ],
            'interest_rate': [
                r'(?:interest\s+rate|rate)[:\s]+(\d+\.?\d*)\s*%?',
                r'rate[:\s]+(\d+\.?\d*)',
            ],
            'occupancy': [
                r'occupancy[:\s]+(\d+\.?\d*)\s*%?',
                r'occupied[:\s]+(\d+\.?\d*)',
            ],
            'square_footage': [
                r'(?:square\s+footage|sq\s*ft|sqft)[:\s]+([\d,]+)',
                r'(\d+)\s*(?:square\s*feet|sq\s*ft)',
            ],
        }
    
    def parse_file(self, file_bytes: bytes, filename: str) -> ExtractionResult:
        """
        Parse any supported file format
        
        Args:
            file_bytes: Raw file bytes
            filename: Original filename (for extension detection)
            
        Returns:
            ExtractionResult with extracted data
        """
        ext = Path(filename).suffix.lower()
        
        try:
            if ext in self.SUPPORTED_FORMATS['pdf']:
                return self._parse_pdf(file_bytes)
            elif ext in self.SUPPORTED_FORMATS['docx']:
                return self._parse_docx(file_bytes)
            elif ext in self.SUPPORTED_FORMATS['text']:
                return self._parse_text(file_bytes)
            elif ext in self.SUPPORTED_FORMATS['email']:
                return self._parse_email(file_bytes, ext)
            elif ext in self.SUPPORTED_FORMATS['image']:
                return self._parse_image(file_bytes)
            elif ext in self.SUPPORTED_FORMATS['spreadsheet']:
                return self._parse_spreadsheet(file_bytes, ext)
            else:
                # Try as text
                return self._parse_text(file_bytes)
        except Exception as e:
            logger.error(f"Error parsing {filename}: {str(e)}")
            return ExtractionResult(
                success=False,
                text="",
                fields={},
                confidence=0.0,
                missing_fields=[],
                error_message=str(e),
                file_type=ext
            )
    
    def _parse_pdf(self, pdf_bytes: bytes) -> ExtractionResult:
        """Parse PDF file"""
        text = ""
        
        # Try pdfplumber first (better for tables)
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")
        
        # Fallback to PyPDF2
        if not text:
            try:
                import PyPDF2
                reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            except Exception as e:
                logger.warning(f"PyPDF2 failed: {e}")
        
        return self._process_extracted_text(text, 'pdf')
    
    def _parse_docx(self, docx_bytes: bytes) -> ExtractionResult:
        """Parse Word document"""
        try:
            import docx
            doc = docx.Document(io.BytesIO(docx_bytes))
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += "\n" + row_text
            
            return self._process_extracted_text(text, 'docx')
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            return ExtractionResult(
                success=False,
                text="",
                fields={},
                confidence=0.0,
                missing_fields=[],
                error_message=f"Could not parse DOCX: {str(e)}",
                file_type='docx'
            )
    
    def _parse_text(self, text_bytes: bytes) -> ExtractionResult:
        """Parse plain text file"""
        try:
            text = text_bytes.decode('utf-8', errors='ignore')
            return self._process_extracted_text(text, 'text')
        except Exception as e:
            return ExtractionResult(
                success=False,
                text="",
                fields={},
                confidence=0.0,
                missing_fields=[],
                error_message=str(e),
                file_type='text'
            )
    
    def _parse_email(self, email_bytes: bytes, ext: str) -> ExtractionResult:
        """Parse email files (.eml, .msg)"""
        text = ""
        
        if ext == '.eml':
            try:
                import eml_parser
                ep = eml_parser.EmlParser()
                parsed = ep.decode_email_bytes(email_bytes)
                
                # Extract body
                if 'body' in parsed:
                    for body in parsed['body']:
                        if 'content' in body:
                            text += body['content'] + "\n"
                
                # Extract headers
                if 'header' in parsed:
                    header = parsed['header']
                    text += f"\nFrom: {header.get('from', '')}\n"
                    text += f"Subject: {header.get('subject', '')}\n"
                    text += f"Date: {header.get('date', '')}\n"
            except Exception as e:
                logger.warning(f"eml_parser failed: {e}")
                # Fallback to simple parsing
                text = email_bytes.decode('utf-8', errors='ignore')
        
        elif ext == '.msg':
            try:
                import extract_msg
                msg = extract_msg.Message(io.BytesIO(email_bytes))
                text = f"Subject: {msg.subject}\n"
                text += f"From: {msg.sender}\n"
                text += f"Date: {msg.date}\n\n"
                text += msg.body
            except Exception as e:
                logger.warning(f"extract_msg failed: {e}")
                text = email_bytes.decode('utf-8', errors='ignore')
        
        return self._process_extracted_text(text, 'email')
    
    def _parse_image(self, image_bytes: bytes) -> ExtractionResult:
        """Parse image with OCR"""
        try:
            from PIL import Image
            import pytesseract
            
            image = Image.open(io.BytesIO(image_bytes))
            
            # OCR the image
            text = pytesseract.image_to_string(image)
            
            return self._process_extracted_text(text, 'image')
        except Exception as e:
            logger.error(f"OCR error: {e}")
            return ExtractionResult(
                success=False,
                text="",
                fields={},
                confidence=0.0,
                missing_fields=[],
                error_message=f"OCR failed: {str(e)}. Make sure tesseract is installed.",
                file_type='image'
            )
    
    def _parse_spreadsheet(self, file_bytes: bytes, ext: str) -> ExtractionResult:
        """Parse spreadsheet files"""
        text = ""
        
        try:
            if ext == '.csv':
                import csv
                csv_text = file_bytes.decode('utf-8', errors='ignore')
                reader = csv.reader(csv_text.splitlines())
                for row in reader:
                    text += " | ".join(row) + "\n"
            
            elif ext in ['.xlsx', '.xls']:
                import pandas as pd
                df = pd.read_excel(io.BytesIO(file_bytes))
                text = df.to_string(index=False)
        except Exception as e:
            logger.error(f"Spreadsheet parsing error: {e}")
            return ExtractionResult(
                success=False,
                text="",
                fields={},
                confidence=0.0,
                missing_fields=[],
                error_message=str(e),
                file_type='spreadsheet'
            )
        
        return self._process_extracted_text(text, 'spreadsheet')
    
    def _process_extracted_text(self, text: str, file_type: str) -> ExtractionResult:
        """Process extracted text and extract fields"""
        if not text:
            return ExtractionResult(
                success=False,
                text="",
                fields={},
                confidence=0.0,
                missing_fields=[],
                error_message="No text extracted from file",
                file_type=file_type
            )
        
        # Extract fields using patterns
        fields = self._extract_fields(text)
        
        # Calculate confidence
        required_fields = ['units', 'address', 'city', 'state', 'zip',
                          'estimated_value', 'purchase_price', 'loan_amount',
                          'credit_score_1', 'credit_score_2', 'credit_score_3']
        found_fields = [f for f in required_fields if f in fields and fields[f]]
        confidence = len(found_fields) / len(required_fields)
        
        missing = [f for f in required_fields if f not in fields or not fields[f]]
        
        return ExtractionResult(
            success=True,
            text=text,
            fields=fields,
            confidence=confidence,
            missing_fields=missing,
            file_type=file_type
        )
    
    def _extract_fields(self, text: str) -> Dict:
        """Extract fields from text using regex patterns"""
        fields = {}
        text_lower = text.lower()
        
        for field_name, patterns in self.patterns.items():
            for pattern in patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    if field_name == 'credit_scores':
                        fields['credit_score_1'] = int(match.group(1))
                        fields['credit_score_2'] = int(match.group(2))
                        fields['credit_score_3'] = int(match.group(3))
                    elif field_name == 'units':
                        fields[field_name] = int(match.group(1))
                    elif field_name in ['estimated_value', 'purchase_price', 'loan_amount', 'points', 'occupancy', 'square_footage']:
                        value = match.group(1).replace(',', '')
                        fields[field_name] = float(value)
                    else:
                        fields[field_name] = match.group(1).strip()
                    break
        
        # Process credit scores if found as separate scores
        if 'credit_score_1' in fields and 'credit_score_2' in fields and 'credit_score_3' in fields:
            scores = sorted([fields['credit_score_1'], fields['credit_score_2'], fields['credit_score_3']])
            fields['credit_score_middle'] = scores[1]
        
        # Try to extract individual credit scores if not found together
        if 'credit_score_1' not in fields:
            credit_matches = re.findall(r'(?:credit\s+score|fico)[:\s]+(\d{3})', text, re.IGNORECASE)
            if len(credit_matches) >= 3:
                fields['credit_score_1'] = int(credit_matches[0])
                fields['credit_score_2'] = int(credit_matches[1])
                fields['credit_score_3'] = int(credit_matches[2])
        
        # Clean up state (convert full name to abbreviation if needed)
        if 'state' in fields:
            state_map = {
                'california': 'CA', 'texas': 'TX', 'florida': 'FL', 'new york': 'NY',
                'illinois': 'IL', 'pennsylvania': 'PA', 'ohio': 'OH', 'georgia': 'GA',
                'north carolina': 'NC', 'michigan': 'MI', 'new jersey': 'NJ',
                'virginia': 'VA', 'washington': 'WA', 'arizona': 'AZ', 'massachusetts': 'MA',
                'tennessee': 'TN', 'indiana': 'IN', 'missouri': 'MO', 'maryland': 'MD',
                'wisconsin': 'WI', 'colorado': 'CO', 'minnesota': 'MN', 'south carolina': 'SC',
                'alabama': 'AL', 'louisiana': 'LA', 'kentucky': 'KY', 'oregon': 'OR',
                'oklahoma': 'OK', 'connecticut': 'CT', 'utah': 'UT', 'iowa': 'IA',
                'nevada': 'NV', 'arkansas': 'AR', 'mississippi': 'MS', 'kansas': 'KS',
                'new mexico': 'NM', 'nebraska': 'NE', 'west virginia': 'WV', 'idaho': 'ID',
                'hawaii': 'HI', 'new hampshire': 'NH', 'maine': 'ME', 'montana': 'MT',
                'rhode island': 'RI', 'delaware': 'DE', 'south dakota': 'SD', 'north dakota': 'ND',
                'alaska': 'AK', 'vermont': 'VT', 'wyoming': 'WY'
            }
            state_val = fields['state'].strip().lower()
            if state_val in state_map:
                fields['state'] = state_map[state_val]
        
        return fields
    
    def get_supported_formats(self) -> List[str]:
        """Return list of supported file extensions"""
        formats = []
        for exts in self.SUPPORTED_FORMATS.values():
            formats.extend(exts)
        return formats


# Convenience function
def parse_loan_file(file_bytes: bytes, filename: str) -> ExtractionResult:
    """Convenience function to parse any supported file"""
    parser = UniversalFileParser()
    return parser.parse_file(file_bytes, filename)
